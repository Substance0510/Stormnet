from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from json import dump as jdump
from json import dumps as jdumps
from fpdf import FPDF
from docx import Document

class DedMoroz:
    def __init__(self, name:str, behavior:bool, gifts=None):
        self.name = name
        self.behavior = behavior
        self.gifts = gifts
        self.out_message = self._print_welcome_message()

        self._print_welcome_message()
        if isinstance(self.gifts, dict):
            self.out_message += self._print_gifts_dict()
        elif isinstance(self.gifts, list):
            self.out_message += self._print_gifts_list()
        elif self.gifts is None and self.behavior:
            self.out_message += 'к сожалению, у Дедушки Мороза не хватило на всех подарков.\n' \
                           'Не расстраивайся, он обязательно что-нибудь придумает ;)\n'
        elif self.gifts is not None and not isinstance(self.gifts, (dict, list)):
            raise Exception('Не поддерживаемый формат подарков.')
        print(self.out_message)

    def _print_welcome_message(self):
        message = f'Здравствуй, мой дорогой друг, {self.name}.\nВ этом году твоё поведение было '
        message += 'хорошим.\n' if self.behavior else 'не очень хорошим.\n'
        message += 'Поэтому, за такое поведение, Дедушка Мороз тебе '
        message += 'принёс:\n' if self.behavior else 'ничего не принёс.\nПостарайся в следующем году.\n'
        return message

    def _print_gifts_dict(self):
        message = ''
        for gift, count in self.gifts.items():
            message += f'> {gift}: {count}\n'
        return message

    def _print_gifts_list(self):
        return '> '+'\n> '.join(self.gifts)

    def _write_file(self):
        file = '/home/anton/PycharmProjects/training/DedMoroz/' + str(self.name) + '.txt'
        with open(file, 'w') as wf:
            wf.write(self.name + '\n' + str(self.behavior) + '\n')
            jdump(self.gifts, wf)

    def write_file_pdf_lab(self):
        pdfmetrics.registerFont(TTFont('DejaVuSerif', 'DejaVuSerif.ttf'))
        file = '/home/anton/PycharmProjects/training/DedMoroz/' + str(self.name) + '.pdf'
        pdf = canvas.Canvas(file)
        pdf.setFont('DejaVuSerif', 20)
        pdf.setTitle('Письмо от Деда Мороза!')
        pdf.drawString(170, 770, 'Письмо от Деда Мороза!')
        pdf.setFont('DejaVuSerif', 15)
        y = 670
        index = 0
        message = self.out_message
        for line in range(len(message)):
            if message[line] == '\n':
                pdf.drawString(20, y, message[index:line])
                index = line + 1
                y -= 20
        pdf.save()

    def write_file_pdf_fpdf(self):
        file = '/home/anton/PycharmProjects/training/DedMoroz/' + str(self.name) + '.pdf'
        pdf = FPDF()
        #pdf.SYSTEM_TTFONTS = '/some/path'
        pdf.add_font('DejaVuSerif', '', 'DejaVuSerif.ttf', uni=True)
        pdf.add_page()
        pdf.set_font("DejaVuSerif", size=15)
        pdf.multi_cell(200, 5, txt=self.out_message)
        pdf.output(file)

    def write_file_docx(self):
        doc = Document()
        doc.add_heading('Письмо от Деда Мороза.')
        doc.add_paragraph(self.out_message)
        doc.save('/home/anton/PycharmProjects/training/DedMoroz/' + str(self.name) + '.docx')


recipient1 = DedMoroz('Anton', True, {'Загородный коттедж со всеми удобствами': 1,
                                      'Новый автомобиль иномарка': 2, 'Квартира в Минске': 3,
                                      'Денежный приз': '5 миллионов долларов США'})
#recipient1.write_file_pdf_lab()
recipient1.write_file_pdf_fpdf()
recipient1.write_file_docx()

recipient2 = DedMoroz(name='Kate', behavior=True, gifts=['Квартира в Минске', 'Новая иномарка',
                                                         'Загородный коттедж', 'Путёвка на Мальдивы.'])

recipient3 = DedMoroz(name='BadBoy', behavior=False)
recipient3.write_file_pdf_lab()

recipient4 = DedMoroz(name='GoodBoy', behavior=True)
recipient4.write_file_pdf_lab()

#recipient5 = DedMoroz(name='TestBoy', behavior=True, gifts='Падаруначки...')